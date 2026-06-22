import re
from datetime import datetime, timedelta
from pathlib import Path

from fastapi import Depends, HTTPException, UploadFile
from sqlalchemy import select, update, func, delete
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload

from app.config import settings
from app.core.milvus_client import get_milvus_client, COLLECTION_NAME
from app.database import get_db
from app.models.device import Device
from app.models.kb import KbDocument, KbChunk
from app.services.embedding_service import EmbeddingService

MAX_TOKENS = 1000


class KbService:
    ALLOWED_EXTENSIONS = {".pdf", ".docx", ".md", ".txt"}

    def __init__(self, db: AsyncSession = Depends(get_db)):
        self.db = db

    async def upload(
            self, file: UploadFile,
            title: str,
            doc_type: str,
            product_model: str | None = None,
            product_series: str | None = None,
            replace_doc_id:int|None=None
    ) -> KbDocument:
        #--- 增量更新时校验旧文档是否存在 ---
        old_doc=None
        if replace_doc_id is not None:
            old_doc=await self.get_document(replace_doc_id)
            if not old_doc:
                raise HTTPException(status_code=404,detail="要替换的文档不存在")
        # ① 校验文件类型
        ext = Path(file.filename).suffix.lower()
        if ext not in self.ALLOWED_EXTENSIONS:
            raise HTTPException(status_code=400, detail=f"不支持的文件类型：{ext}")
        # ② 保存到磁盘
        file_path = await self._save_to_disk(file)
        # ③ 按格式读取文字内容
        if ext == ".pdf":
            content = self._read_pdf(file_path)
        elif ext == ".docx":
            content = self._read_docx(file_path)
        else:
            content = self._read_text(file_path)

        if not content.strip():
            raise HTTPException(status_code=400, detail="文件内容为空")

        # ④ 切块
        chunks = self._structure_aware_chunk(content)

        # ⑤ 创建或复用文档记录
        if old_doc:
            doc=old_doc
        else:
            doc = KbDocument(
                title=title or file.filename,
                doc_type=doc_type,
                file_path=file_path,
                product_model=product_model,
                product_series=product_series
            )
            self.db.add(doc)
            await self.db.flush()

        # ⑥ 保存分块到mysql
        chunk_objects=[]
        for i, c in enumerate(chunks):
            chunk=KbChunk(
                document_id=doc.id,
                chunk_index=i,
                content=c["content"],
                chunk_type=c["type"],
                parent_title=c.get("parent_title"),
                token_count=c.get("token_count")
            )
            self.db.add(chunk)
            chunk_objects.append(chunk)

        await self.db.flush()

        # ⑦ 向量化 + 存入 Milvus
        new_milvus_ids=[]
        try:
            texts=[chunk.content for chunk in chunk_objects]
            embeddings=EmbeddingService.embed_batch(texts)

            milvus=get_milvus_client()
            milvus_data=[]
            for chunk,emb in zip(chunk_objects,embeddings):
                milvus_data.append({
                    "id":chunk.id,
                    "vector":emb,
                    "chunk_id":chunk.id,
                    "content":chunk.content,
                    "content_type":chunk.chunk_type,
                    "parent_title":chunk.parent_title or "",
                    "document_id":doc.id,
                    "product_series": doc.product_series or "",
                    "doc_version": str(doc.version),
                })
                chunk.milvus_id=chunk.id
                new_milvus_ids.append(chunk.id)

            milvus.insert(collection_name=COLLECTION_NAME,data=milvus_data)
            milvus.flush()
        except Exception as e:
            raise HTTPException(status_code=500,detail=f"Milvus写入失败：{e}")

        # ⑧ 增量更新：验证新向量 → 删旧数据
        if old_doc: #如果旧文档存在
            #查询 Milvus 中是否已经存在与当前文档关联的新向量数据
            verify=milvus.query(
                collection_name=COLLECTION_NAME,
                expr=f"document_id=={doc.id}",
                limit=1
            )
            if not verify:
                #如果查不到新数据，说明在执行这段代码之前，新向量的写入操作失败了。
                # 此时需要执行回滚操作：删除可能部分写入的残缺新向量（new_milvus_ids）
                milvus.delete(
                    collection_name=COLLECTION_NAME,
                    expr=f"id in {new_milvus_ids}"
                )
                raise HTTPException(status_code=500,detail="新向量写入失败，已回滚")
            # 删除旧 chunk（MySQL + Milvus）
            #TODO:存在事务一致性缺失，如果 MySQL 删除失败，Milvus 中的向量已经删除，
            # 但 MySQL 中的记录还在，会导致脏数据（用户查不到向量，但数据库里还有记录）
            old_ids=[c.id for c in old_doc.chunks]
            if old_ids:
                milvus.delete(
                    collection_name=COLLECTION_NAME,
                    expr=f"document_id=={doc.id} and id not in {new_milvus_ids}"
                )
                await self.db.execute(
                    delete(KbChunk).where(KbChunk.id.in_(old_ids))
                )
            doc.file_path=file_path
            doc.version=doc.version+1
            doc.status="active"

        doc.chunk_count = len(chunks)
        doc.updated_at=datetime.now()
        await self.db.flush()
        await self.db.refresh(doc)
        return doc

    # 文件保存
    @staticmethod
    async def _save_to_disk(file: UploadFile) -> str:
        upload_dir = Path(settings.UPLOAD_DIR)
        #如果不存在则递归创建
        upload_dir.mkdir(parents=True, exist_ok=True)

        # re 是 Python 的标准库模块，代表正则表达式1。它提供了一套强大的工具，用于字符串的模式匹配、搜索、替换和解析等操作
        #匹配不是字母、数字、下划线(\w)、连字符(-)和点号(.)的任何字符，并将其替换为下划线 _
        #目的：防止路径遍历攻击。如果恶意用户上传的文件名为 ../../etc/passwd，经过此正则处理后，会变成 .._.._etc_passwd，从而避免了文件被写到预期目录之外的危险
        safe_name = re.sub(r"[^\w\-.]", "_", file.filename)
        #使用/ 运算符拼接路径
        file_path = upload_dir / safe_name

        # 如果目标路径已经存在同名文件，代码会自动在文件名后添加带括号的序号
        if file_path.exists():
            stem, suffix = file_path.stem, file_path.suffix
            counter = 1
            while file_path.exists():
                file_path = upload_dir / f"{stem}({counter}){suffix}"
                counter += 1

        file_path.write_bytes(await file.read())
        return str(file_path)

    # 三种格式读取
    @staticmethod
    def _read_text(path: str) -> str:
        try:
            return Path(path).read_text(encoding="utf-8")
        except UnicodeDecodeError:
            return Path(path).read_text(encoding="gbk")

    @staticmethod
    def _read_pdf(path: str) -> str:
        #即 PyMuPDF
        import fitz
        doc = fitz.open(path)
        all_parts = []

        for page in doc:
            blocks = page.get_text("dict")["blocks"]
            for block in blocks:
                # type=0：文本块
                if block["type"] == 0:
                    text = " ".join(
                        span["text"] for line in block["lines"]
                        for span in line["spans"]
                    ).strip()
                    if len(text) < 5:
                        continue
                        #检测是否为表格行（按坐标判断：同一行有多个对齐的文本块）
                    all_parts.append(text)
                # type=1：图片块 → 跳过（不做 OCR），只记录占位
                elif block["type"] == 1:
                    all_parts.append("[图片]")

        doc.close()
        return "\n\n".join(all_parts)

    @staticmethod
    def _read_docx(path: str) -> str:
        from docx import Document
        doc = Document(path)
        """
        使用列表推导式遍历文档的所有段落（doc.paragraphs），提取段落文本。
        if p.text.strip() 的作用是过滤掉空段落或仅包含空白字符的段落。
        """
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            """
            对于每个表格，遍历其每一行（table.rows），再遍历每行中的每个单元格（row.cells），提取单元格文本并去除首尾空白。
            然后使用 | 作为分隔符，将同一行的单元格文本拼接起来，形成一个类似 Markdown 表格格式的字符串
            """
            rows = ["|".join(cell.text.strip() for cell in row.cells) for row in table.rows]
            parts.append("\n".join(rows))
        #将 parts 列表中的所有段落文本和表格文本用两个换行符（\n\n）拼接，模拟文档中段落与段落、段落与表格之间的空行分隔效果
        return "\n\n".join(parts)

    # 结构感知分块
    @staticmethod
    def _structure_aware_chunk(content: str) -> list[dict]:
        # 第一步：按H2/H3标题切分章节
        sections = KbService._split_by_headings(content)

        # 第二步：每个章节内部，分离表格/步骤/普通段落
        candidates = []
        for section_title, section_body in sections:
            candidates.extend(KbService._extract_candidates(section_body, section_title))

        # 第三步：超长块二次拆分
        result = []
        for c in candidates:
            if c["token_count"] <= MAX_TOKENS:
                result.append(c)
            else:
                result.extend(KbService._split_long(c))
        return result

    @staticmethod
    def _split_by_headings(content: str) -> list[tuple[str, str]]:
        pattern = r"^#{2,3}\s+.*$"
        lines = content.split("\n")
        sections = []
        current_title = ""
        current_body: list[str] = []

        """
        当遇到一个新的符合正则的标题行时，代码会先将之前缓存的正文拼接起来，
        与旧标题组成元组存入 sections 列表；然后清空正文缓存，并更新当前标题为新标题
        """
        for line in lines:
            if re.match(pattern, line.strip()):
                if current_body:
                    body = "\n".join(current_body).strip()
                    if body:
                        sections.append((current_title, body))
                    current_body = []
                current_title = line.strip().lstrip("#").strip()
            else:
                current_body.append(line)

        if current_body:
            body = "\n".join(current_body).strip()
            if body:
                sections.append((current_title, body))

        return sections

    @staticmethod
    def _extract_candidates(body: str, parent_title: str) -> list[dict]:
        candidates = []
        paragraphs = body.split("\n\n")
        i = 0

        while i < len(paragraphs):
            para = paragraphs[i].strip()
            if not para:
                i += 1
                continue

            # 检测表格
            if KbService._is_table(para):
                table_lines = [para]
                i += 1
                #继续检测后续段落，直到遇到非表格段落为止
                while i < len(paragraphs) and KbService._is_table(paragraphs[i].strip()):
                    table_lines.append(paragraphs[i].strip())
                    i += 1
                content = "\n".join(table_lines)
                prefix = f"[{parent_title}]\n" if parent_title else ""
                candidates.append({
                    "content": prefix + content,
                    "type": "table",
                    "parent_title": parent_title,
                    "token_count": KbService._count_tokens(content)
                })
                continue

            # 检测有序步骤
            if re.match(r"^\d+[.)、]\s", para):
                step_lines = [para]
                i += 1
                #将连续的有序步骤段落合并拼接
                while i < len(paragraphs):
                    nxt = paragraphs[i].strip()
                    if re.match(r"^\d+[.)、]\s", nxt):
                        step_lines.append(nxt)
                        i += 1
                    else:
                        break
                content = "\n".join(step_lines)
                candidates.append({
                    "content": content,
                    "type": "step",
                    "parent_title": parent_title,
                    "token_count": KbService._count_tokens(content)
                })
                continue

            # 普通段落
            candidates.append({
                "content": para,
                "type": "paragraph",
                "parent_title": parent_title,
                "token_count": KbService._count_tokens(para)
            })
            i += 1

        return candidates

    @staticmethod
    def _is_table(line: str) -> bool:
        """"判断是否为表格行：以｜开头且至少含两个｜"""
        return line.startswith("|") and line.count("|") >= 2

    @staticmethod
    def _split_long(chunk: dict) -> list[dict]:
        """超长块二次拆分：按段落切，每段头部附加 '[父标题]' 前缀"""
        title = chunk.get("parent_title", "")
        prefix = f"[{title}]\n" if title else ""
        paragraphs = chunk["content"].split("\n\n")
        result = []
        cur = ""

        for p in paragraphs:
            p = p.strip()
            if not p:
                continue
            if KbService._count_tokens(cur + "\n\n" + p) > MAX_TOKENS and cur:
                result.append({
                    "content": prefix + cur,
                    "type": chunk["type"],
                    "parent_title": title,
                    "token_count": KbService._count_tokens(cur)
                })
                cur = p
            else:
                cur = cur + "\n\n" + p if cur else p

        if cur:
            result.append({
                "content": prefix + cur,
                "type": chunk["type"],
                "parent_title": title,
                "token_count": KbService._count_tokens(cur)
            })
        return result

    # 计数token工具
    @staticmethod
    def _count_tokens(text: str) -> int:
        # 匹配文本中的所有中文字符（该 Unicode 范围涵盖了绝大多数常用汉字）
        cn = len(re.findall(r"[\u4e00-\u9fff]", text))
        # 通常 1 个汉字大约会被编码为 1.5 到 2 个 Token
        # 1 个英文单词大约是 1 个 Token，而英文单词平均长度约为 4 个字符，因此这里除以 4，相当于认为 4 个非中文字符 ≈ 1 个 Token
        return int(cn / 1.5 + (len(text) - cn) / 4)

    # 查询
    async def get_document(self, doc_id: int) -> KbDocument | None:
        r = await self.db.execute(
            select(KbDocument)
            .where(KbDocument.id == doc_id)
            .options(selectinload(KbDocument.chunks))
        )
        return r.scalar_one_or_none()

    async def list_documents(self, skip: int = 0, limit: int = 20) -> list[KbDocument]:
        r = await self.db.execute(
            select(KbDocument)
            .order_by(KbDocument.created_at.desc())
            .offset(skip)
            .limit(limit)
        )
        return list(r.scalars().all())

    async def search_chunks(self, query: str, top_k: int = 5) -> list[dict]:
        # ① 把用户问题转为向量
        query_embedding = EmbeddingService.embed_single(query)

        # ② 在 Milvus 中搜索（只拿 chunk_id + document_id）
        milvus = get_milvus_client()
        results = milvus.search(
            collection_name=COLLECTION_NAME,
            data=[query_embedding],
            limit=top_k * 3,
            output_fields=["chunk_id", "document_id"]
        )

        # ③ 从 MySQL 查分块完整内容
        chunk_ids = list({hit["id"] for hit in results[0]})
        chunk_map = {}
        if chunk_ids:
            r = await self.db.execute(
                select(KbChunk).where(KbChunk.id.in_(chunk_ids))
            )
            chunk_map = {c.id: c for c in r.scalars().all()}

        # 收集所有 document_id，批量查状态
        doc_ids = list({
            hit.get("entity", {}).get("document_id", 0)
            for hit in results[0]
        })
        doc_status_map = {}
        if doc_ids:
            r = await self.db.execute(
                select(KbDocument.id, KbDocument.status).where(
                    KbDocument.id.in_(doc_ids)
                )
            )
            doc_status_map = {row[0]: row[1] for row in r.all()}

        # 过滤 + 降权 + 填入内容
        hits = []
        for hit in results[0]:
            entity = hit.get("entity", {})
            doc_id = entity.get("document_id", 0)
            status = doc_status_map.get(doc_id, "active")

            if status == "archived":
                continue

            chunk = chunk_map.get(hit["id"])
            if not chunk:
                continue

            score = hit["distance"]
            if status == "expired":
                score *= 0.3

            hits.append({
                "chunk_id": hit["id"],
                "content": chunk.content,
                "chunk_type": chunk.chunk_type,
                "parent_title": chunk.parent_title,
                "document_id": doc_id,
                "score": score
            })

        # 重新按 score 排序，截取 top_k
        hits.sort(key=lambda x: x["score"], reverse=True)
        return hits[:top_k]

    #生命周期管理
    async def update_status(self,doc_id:int,new_status:str)->KbDocument:
        allowed={"active","review_due","expired","archived"}
        if new_status not in allowed:
            raise HTTPException(status_code=400,detail=f"无效状态")

        doc=await self.get_document(doc_id)
        if not doc:
            raise HTTPException(status_code=404,detail="文档不存在")
        doc.status=new_status
        await self.db.flush()
        await self.db.refresh(doc)
        return doc

    async def renew_document(self,doc_id:int)->KbDocument:
        doc=await self.get_document(doc_id)
        if not doc:
            raise HTTPException(status_code=404,detail="文档不存在")

        doc.status="active"
        doc.updated_at=datetime.now()
        await self.db.flush()
        await self.db.refresh(doc)
        return doc

    async def scan_expired(self)->dict:
        """
        每日定时扫描：
        - active 超过 90 天未更新 → review_due
        - expired 超过 180 天 → archived
        """
        now=datetime.now()
        threshold_90d=now-timedelta(days=90)
        threshold_180d=now-timedelta(days=180)

        r1=await self.db.execute(
            update(KbDocument)
            .where(KbDocument.status=="active",KbDocument.updated_at<threshold_90d)
            .values(status="review_due")
        )
        marked_review_due=r1.rowcount

        r2=await self.db.execute(
            update(KbDocument)
            .where(KbDocument.status=="expired",KbDocument.updated_at<threshold_180d)
            .values(status="archived")
        )
        archived_expired=r2.rowcount

        await self.db.flush()
        total=await self.db.scalar(select(func.count()).select_from(KbDocument))

        return {
            "scanned":total or 0,
            "marked_review_due":marked_review_due,
            "archived_expired":archived_expired
        }

    async def mark_by_device_status(self,model_number:str)->int:
        """设备退市/固件更新时回调，关联该型号的 active 文档标记为 review_due"""
        r=await self.db.execute(
            update(KbDocument)
            .where(KbDocument.product_model==model_number,KbDocument.status=="active")
            .values(status="review_due")
        )
        await self.db.flush()
        return r.rowcount

    async def get_top_referenced(self,limit:int=10)->list[KbDocument]:
        """高频引用统计"""
        r=await self.db.execute(
            select(KbDocument)
            .order_by(KbDocument.reference_count.desc())
            .limit(limit)
        )
        return list(r.scalars().all())



